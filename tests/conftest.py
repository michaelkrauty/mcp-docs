"""Shared pytest fixtures for mcp-docs tests."""

import tempfile
from pathlib import Path
from uuid import uuid4

import pytest


def qdrant_available() -> bool:
    """Check if Qdrant is running."""
    import httpx

    try:
        response = httpx.get("http://localhost:6333/collections", timeout=2.0)
        return response.status_code == 200
    except Exception:
        return False


def embedding_available() -> bool:
    """Check if embedding service is running."""
    import httpx

    try:
        response = httpx.get("http://localhost:8080/health", timeout=2.0)
        return response.status_code == 200
    except Exception:
        return False


requires_qdrant = pytest.mark.skipif(
    not qdrant_available(),
    reason="Qdrant not available at localhost:6333",
)

requires_embedding = pytest.mark.skipif(
    not embedding_available(),
    reason="Embedding service not available at localhost:8080",
)

requires_services = pytest.mark.skipif(
    not (qdrant_available() and embedding_available()),
    reason="Qdrant and/or embedding service not available",
)


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def test_collection_name():
    """Generate unique test collection name."""
    return f"docs_test_{uuid4().hex[:12]}"


@pytest.fixture
async def qdrant_storage():
    """Create QdrantStorage instance for testing."""
    from vector_core import QdrantStorage

    storage = QdrantStorage(url="http://localhost:6333")
    yield storage
    await storage.close()


@pytest.fixture
async def test_collection(qdrant_storage, test_collection_name):
    """Create a test collection and clean up after."""
    await qdrant_storage.create_collection(test_collection_name)
    yield test_collection_name
    try:
        await qdrant_storage.delete_collection(test_collection_name)
    except Exception:
        pass


@pytest.fixture
def sample_pdf(temp_dir: Path) -> Path:
    """Create a sample PDF for testing (minimal valid PDF)."""
    file_path = temp_dir / "sample.pdf"
    # Minimal PDF structure
    file_path.write_bytes(
        b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj 2 0 "
        b"obj<</Type/Pages/Count 0/Kids[]>>endobj xref\n0 3\n0000000000 "
        b"65535 f\n0000000009 00000 n\n0000000052 00000 n\ntrailer<</Size 3"
        b"/Root 1 0 R>>startxref\n101\n%%EOF"
    )
    return file_path


@pytest.fixture
def sample_docx(temp_dir: Path) -> Path:
    """Create a sample DOCX for testing."""
    from docx import Document

    file_path = temp_dir / "sample.docx"
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph with some content.")
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def sample_text(temp_dir: Path) -> Path:
    """Create a sample text file for testing."""
    file_path = temp_dir / "sample.txt"
    file_path.write_text(
        "This is a sample text file.\n\n"
        "It contains multiple paragraphs.\n\n"
        "This is useful for testing document extraction."
    )
    return file_path


@pytest.fixture
def sample_markdown(temp_dir: Path) -> Path:
    """Create a sample markdown file for testing."""
    file_path = temp_dir / "sample.md"
    file_path.write_text(
        "# Test Document\n\n"
        "This is a **markdown** file.\n\n"
        "## Section 1\n\n"
        "Some content here.\n\n"
        "## Section 2\n\n"
        "More content here."
    )
    return file_path


@pytest.fixture
def sample_dense_vector():
    """Sample dense vector for testing."""
    from mcp_docs.settings import settings

    return [0.1] * settings.embedding_dim


@pytest.fixture
def sample_sparse_vector():
    """Sample sparse vector for testing."""
    from vector_core.embeddings.sparse import SparseVector

    return SparseVector(
        indices=[0, 5, 10, 15, 100],
        values=[0.5, 0.3, 0.2, 0.15, 0.1],
    )


@pytest.fixture
def document_store(temp_dir: Path):
    """Create a temporary DocumentStore for testing."""
    from mcp_docs.storage.database import DocumentStore

    db_path = temp_dir / "test_docs.db"
    store = DocumentStore(db_path=db_path)
    yield store
    store.close()


@pytest.fixture
async def search_engine(test_collection_name, qdrant_storage):
    """Create DocumentSearchEngine for testing with cleanup."""
    from mcp_docs.search import DocumentSearchEngine

    # Create engine with test collection
    engine = DocumentSearchEngine(collection_name=test_collection_name)
    yield engine
    await engine.close()


@pytest.fixture
def document_root(temp_dir: Path) -> Path:
    """Create a document root directory with sample files."""
    root = temp_dir / "documents"
    root.mkdir()

    # Create some sample files
    (root / "readme.txt").write_text("This is a readme file.")
    (root / "notes.md").write_text("# Notes\n\nSome notes here.")

    subdir = root / "reports"
    subdir.mkdir()
    (subdir / "report.txt").write_text("Quarterly report content.")

    return root
